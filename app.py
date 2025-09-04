import os
from io import BytesIO
from dotenv import load_dotenv
from openai import OpenAI
import streamlit as st
import docx
import PyPDF2
import requests
from bs4 import BeautifulSoup
import pathlib
import numpy as np

st.set_page_config(page_title="AI CV Upgrader", page_icon="🚀", layout="centered")

# ---------------- Last File / LinkedIn Memory ----------------
LAST_CV_FILE = "last_cv_path.txt"
LAST_LINKEDIN_FILE = "last_linkedin.txt"

def save_last_cv_path(path: str):
    with open(LAST_CV_FILE, "w") as f:
        f.write(path)

def load_last_cv_path() -> str:
    if pathlib.Path(LAST_CV_FILE).exists():
        with open(LAST_CV_FILE, "r") as f:
            return f.read().strip()
    return ""

def save_last_linkedin(linkedin_url: str):
    with open(LAST_LINKEDIN_FILE, "w") as f:
        f.write(linkedin_url)

def load_last_linkedin() -> str:
    if pathlib.Path(LAST_LINKEDIN_FILE).exists():
        with open(LAST_LINKEDIN_FILE, "r") as f:
            return f.read().strip()
    return ""

# ---------------- File Readers ----------------
def read_docx_bytes(file_bytes: bytes) -> str:
    bio = BytesIO(file_bytes)
    document = docx.Document(bio)
    return "\n".join(p.text for p in document.paragraphs)

def read_pdf_bytes(file_bytes: bytes) -> str:
    bio = BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(bio)
    text = ""
    for page in reader.pages:
        text += (page.extract_text() or "") + "\n"
    return text.strip()

# ---------------- Scraper ----------------
def scrape_text_from_url(url: str, timeout: int = 15) -> str:
    try:
        resp = requests.get(url, timeout=timeout, headers={"User-Agent":"Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        return soup.get_text(separator="\n", strip=True)[:20000]
    except Exception:
        return ""

# ---------------- DOCX Saver ----------------
def save_text_to_docx_bytes(text: str) -> bytes:
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------------- OpenAI Client ----------------
def get_client() -> OpenAI:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("OPENAI_API_KEY not set in .env")
        st.stop()
    return OpenAI(api_key=api_key)

# ---------------- Prompts ----------------
SYSTEM_PROMPT_CV = """You are a senior career coach and expert resume writer.
Tailor the candidate's CV to the job and optional LinkedIn info. ATS-friendly, professional tone, concise bullets.
"""

USER_PROMPT_CV = """Candidate CV:
{cv_text}

LinkedIn Info:
{linkedin_text}

Target Job Description:
{job_text}

Task: Rewrite the CV to best match the target role. Return only upgraded CV content.
"""

SYSTEM_PROMPT_COVER = """You are a senior career coach and expert resume writer.
Write a tailored, compelling one-page cover letter aligned to the job description and candidate's background. 250-400 words, professional and concise.
"""

USER_PROMPT_COVER = """Candidate CV:
{cv_text}

LinkedIn Info:
{linkedin_text}

Target Job Description:
{job_text}

Task: Produce a polished cover letter addressed generically (Dear Hiring Manager,) in plain text.
"""

# ---------------- OpenAI Functions ----------------
def generate_with_openai(client: OpenAI, system_prompt: str, user_prompt: str, temperature: float = 0.25, model: str = "gpt-4o") -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}],
        temperature=temperature,
    )
    return resp.choices[0].message.content.strip()

def compute_match_percentage(client: OpenAI, candidate_text: str, job_text: str) -> float:
    emb_candidate_resp = client.embeddings.create(input=candidate_text, model="text-embedding-3-small")
    emb_job_resp = client.embeddings.create(input=job_text, model="text-embedding-3-small")
    emb_candidate = emb_candidate_resp.data[0].embedding
    emb_job = emb_job_resp.data[0].embedding
    a = np.array(emb_candidate)
    b = np.array(emb_job)
    similarity = np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
    return round(similarity * 100, 2)

# ---------------- Streamlit UI ----------------
st.title("🚀 AI CV Upgrader")

# Load last LinkedIn
last_linkedin = load_last_linkedin()

# User Inputs
cv_file = st.file_uploader("Upload your CV (.docx or .pdf)", type=["docx","pdf"])
job_text = st.text_area("Paste job description here")
job_url = st.text_input("Or provide job URL (optional)")
linkedin_url = st.text_input("LinkedIn profile URL (optional)", value=last_linkedin)
cover_letter = st.checkbox("Generate cover letter too")

# Load last CV if no new one uploaded
if not cv_file:
    last_path = load_last_cv_path()
    if last_path and pathlib.Path(last_path).exists():
        cv_file = open(last_path, "rb")
        st.info(f"Using last uploaded CV: {last_path}")

# Generate CV & Cover Letter
if st.button("Generate Upgraded CV"):
    if not cv_file:
        st.error("Please upload a CV or make sure last CV exists.")
        st.stop()
    if not (job_text or job_url):
        st.error("Please provide job description text or job URL.")
        st.stop()

    client = get_client()

    # Read CV
    if hasattr(cv_file, "read"):
        cv_bytes = cv_file.read()
        ext = pathlib.Path(getattr(cv_file, "name", "cv.docx")).suffix.lower()
        if ext == ".docx":
            cv_text_content = read_docx_bytes(cv_bytes)
        elif ext == ".pdf":
            cv_text_content = read_pdf_bytes(cv_bytes)
        else:
            st.error("Unsupported CV format")
            st.stop()
    else:  # file from path
        with open(cv_file, "rb") as f:
            cv_bytes = f.read()
            if cv_file.lower().endswith(".docx"):
                cv_text_content = read_docx_bytes(cv_bytes)
            else:
                cv_text_content = read_pdf_bytes(cv_bytes)

    # Save last CV path
    if hasattr(cv_file, "name"):
        save_last_cv_path(cv_file.name)

    # Save LinkedIn URL
    if linkedin_url:
        save_last_linkedin(linkedin_url)

    # Get job text
    if job_url:
        job_text_scraped = scrape_text_from_url(job_url)
        if job_text_scraped.strip():
            job_text = job_text_scraped

    # Get LinkedIn info
    linkedin_text_content = ""
    if linkedin_url:
        linkedin_text_content = scrape_text_from_url(linkedin_url) or linkedin_url

    # Generate upgraded CV
    st.info("Generating upgraded CV...")
    upgraded_cv = generate_with_openai(client, SYSTEM_PROMPT_CV,
        USER_PROMPT_CV.format(cv_text=cv_text_content, linkedin_text=linkedin_text_content, job_text=job_text),
        temperature=0.2)

    # Compute match percentage
    match_percent = compute_match_percentage(client, cv_text_content + "\n" + linkedin_text_content, job_text)
    st.metric("CV Match with Job Description", f"{match_percent}%")
    st.progress(int(match_percent))

    # Download CV
    cv_docx_bytes = save_text_to_docx_bytes(upgraded_cv)
    st.download_button("Download Upgraded CV", cv_docx_bytes,
                       file_name="Upgraded_CV.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Generate cover letter if requested
    if cover_letter:
        st.info("Generating cover letter...")
        cover = generate_with_openai(client, SYSTEM_PROMPT_COVER,
            USER_PROMPT_COVER.format(cv_text=cv_text_content, linkedin_text=linkedin_text_content, job_text=job_text),
            temperature=0.25)
        cover_bytes = save_text_to_docx_bytes(cover)
        st.download_button("Download Cover Letter", cover_bytes,
                           file_name="Cover_Letter.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
