import os
import argparse
from typing import Optional
from dotenv import load_dotenv
from openai import OpenAI
import docx
import PyPDF2
import requests
from bs4 import BeautifulSoup

# ------------------ Last CV Memory ------------------
LAST_CV_FILE = "last_cv.txt"

def save_last_cv(cv_path: str):
    with open(LAST_CV_FILE, "w") as f:
        f.write(cv_path)

def load_last_cv() -> str:
    if os.path.exists(LAST_CV_FILE):
        with open(LAST_CV_FILE, "r") as f:
            return f.read().strip()
    return ""

# ------------------ File Readers ------------------
def read_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text() or ""
            text += t + "\n"
    return text.strip()

def read_cv(file_path: str) -> str:
    if file_path.lower().endswith(".docx"):
        return read_docx(file_path)
    elif file_path.lower().endswith(".pdf"):
        return read_pdf(file_path)
    else:
        raise ValueError("Unsupported CV format. Please use .docx or .pdf")

# ------------------ Scraping ------------------
def scrape_text_from_url(url: str, timeout: int = 15) -> str:
    try:
        resp = requests.get(url, timeout=timeout, headers={"User-Agent":"Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        return text[:20000]  # limit size
    except Exception:
        return ""

# ------------------ DOCX Saver ------------------
def save_to_docx(text: str, output_path: str) -> None:
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)

# ------------------ OpenAI ------------------
def get_client() -> OpenAI:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set in .env")
    return OpenAI(api_key=api_key)

SYSTEM_PROMPT_CV = """You are a senior career coach and expert resume writer.
Tailor the candidate's CV to the provided job description and optional LinkedIn info.
Keep ATS-friendly formatting, professional tone, concise bullets, standard headers.
Do NOT invent employment; sensible rephrasing is OK.
"""

USER_PROMPT_CV = """Candidate CV:
{cv_text}

LinkedIn Info:
{linkedin_text}

Target Job Description:
{job_text}

Task: Rewrite the CV to best match the target role. Emphasize relevant experience, achievements, and keywords. Return only upgraded CV content.
"""

SYSTEM_PROMPT_COVER = """You are a senior career coach and expert resume writer.
Write a tailored, compelling one-page cover letter aligned to the job description and candidate's background. 250-400 words, professional and concise.
"""

USER_PROMPT_COVER = """Candidate CV:
{cv_text}

LinkedIn Info:
{linkedin_text}

Target Job Description:
{job_text}

Task: Produce a polished cover letter addressed generically (Dear Hiring Manager,) in plain text.
"""

def generate_with_openai(client: OpenAI, system_prompt: str, user_prompt: str, temperature: float = 0.25, model: str = "gpt-4o") -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}],
        temperature=temperature,
    )
    return resp.choices[0].message.content.strip()

def generate_upgraded_cv(cv_text: str, job_text: str, linkedin_text: str, client: OpenAI, model: str = "gpt-4o") -> str:
    user = USER_PROMPT_CV.format(cv_text=cv_text, job_text=job_text, linkedin_text=linkedin_text)
    return generate_with_openai(client, SYSTEM_PROMPT_CV, user, temperature=0.2, model=model)

def generate_cover_letter(cv_text: str, job_text: str, linkedin_text: str, client: OpenAI, model: str = "gpt-4o") -> str:
    user = USER_PROMPT_COVER.format(cv_text=cv_text, job_text=job_text, linkedin_text=linkedin_text)
    return generate_with_openai(client, SYSTEM_PROMPT_COVER, user, SYSTEM_PROMPT_COVER, temperature=0.25, model=model)

# ------------------ CLI ------------------
def main():
    parser = argparse.ArgumentParser(description="AI CV Upgrader")
    parser.add_argument("--cv", help="Path to your CV (.docx or .pdf). If omitted, last CV is used.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--job-url", help="URL to job description")
    group.add_argument("--job-text", help="Paste job description text directly")
    parser.add_argument("--linkedin", help="Your LinkedIn URL or text (optional)")
    parser.add_argument("--cover-letter", action="store_true", help="Also generate a tailored cover letter")
    parser.add_argument("--model", default="gpt-4o", help="OpenAI model (default: gpt-4o)")
    parser.add_argument("--cv-out", default="Upgraded_CV.docx", help="Output path for upgraded CV DOCX")
    parser.add_argument("--cover-out", default="Cover_Letter.docx", help="Output path for cover letter DOCX")
    args = parser.parse_args()

    cv_path = args.cv or load_last_cv()
    if not cv_path:
        raise ValueError("No CV provided and no previous CV found. Please provide a CV.")
    else:
        save_last_cv(cv_path)

    cv_text = read_cv(cv_path)
    linkedin_text = ""
    if args.linkedin:
        linkedin_text = scrape_text_from_url(args.linkedin) or args.linkedin

    if args.job_url:
        job_text = scrape_text_from_url(args.job_url)
        if not job_text.strip():
            print("[!] Scrape returned empty text. Consider using --job-text instead.")
    else:
        job_text = args.job_text

    client = get_client()

    print("[*] Generating upgraded CV...")
    upgraded_cv = generate_upgraded_cv(cv_text, job_text, linkedin_text, client, model=args.model)
    save_to_docx(upgraded_cv, args.cv_out)
    print(f"[✓] Saved upgraded CV -> {args.cv_out}")

    if args.cover_letter:
        print("[*] Generating cover letter...")
        cover = generate_cover_letter(cv_text, job_text, linkedin_text, client, model=args.model)
        save_to_docx(cover, args.cover_out)
        print(f"[✓] Saved cover letter -> {args.cover_out}")

if __name__ == "__main__":
    main()
